# -*- coding: utf-8 -*-

# lstm
from __future__ import print_function
from lexrankr import LexRank
from keras.preprocessing import sequence
from sklearn.feature_extraction.text import TfidfVectorizer  # tfidf,sklearn설정에서 깔아주기
from collections import defaultdict

# docx파일
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

from konlpy.tag import Twitter
import pickle

import sys

reload(sys)
sys.setdefaultencoding('utf-8')


# os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = '/home/ubuntu/smb/HaniumProject-f48c4e461bf5.json'


def txt_to_docx(text_path, filename):
    twitter = Twitter()

    # tokenizer 불러오기
    with open("./tokenizer.pickle", 'rb') as fp:
        RH_tokenizer = pickle.load(fp)
    # LSTM 모델 불러오기
    from keras.models import load_model

    model = load_model('./RH_lstm2.model')  # lstm 모델 경로

    # STT된 텍스트 파일 경로
    text = open(text_path, 'r', encoding='utf-8')
    sen_tag = []  # tagging job
    verb_1 = ['다', '요', '죠', '까', '니']

    dataset = text.read()
    morph = twitter.pos(dataset)
    sen_tag.append(morph)
    verb = []
    for s1 in sen_tag:
        for word, tag in s1:
            if tag == 'Verb':
                if word[-1] in verb_1:
                    verb.append(word)
    verb = list(set(verb))
    for i in verb:
        dataset = dataset.replace(i, i + '.\n')

    # 형태소 분석기로 부족한 동사들 dotword로 충전
    dotword = ['립니다', '까요', '니다', '었다', '해요', '혔다', '였다', '든요', '거든요', '아요', '습니다', '합니다', '입니다', '랍니다', '씁시다', '합시다',
               '하죠', '보죠', '이죠', '렇죠', '시죠', '었고요', '였고요', '니고요', '렇고요', '니까요', '는데>요', '니까', '이에요', '잖아요', "고요", "구요",
               "게요"]
    for i in dotword:
        text = dataset.replace(i + ' ', i + '.\n')

    line = text.readlines()

    # 문단(5문장씩) 나누기
    text_li = []
    i = 0

    while (i + 4 < len(line)):
        line4 = ''.join(line[i:i + 5]).replace('\n', ' ')

        text_li.append(line4)
        i += 5

    if (i < len(line)) and (i + 4 >= len(line)):
        line4 = ''.join(line[i:]).replace('\n', ' ')
        # text_li.append(line4)
        text_li[-1] += line4
        i += 5

    # output >> 문단별 카테고리 분류
    sequences_6m = RH_tokenizer.texts_to_sequences(text_li)
    sequences_matrix_6m = sequence.pad_sequences(sequences_6m, maxlen=500)
    output = model.predict_classes(sequences_matrix_6m)

    # 분류해서 딕셔너리에 카테고리 같은문단끼리 저장

    dict = {}
    for i in range(len(output)):
        if output[i] in dict:
            dict[output[i]].append(text_li[i])
        else:
            dict[output[i]] = [text_li[i]]

    # v_li >> 묶어진 문단v를 v_li에 저장
    v_li = []
    for v in dict.values():
        v_li.append(v)

    ########키워드 추출#########

    def rmmult(sort):
        lili = []
        for i in range(len(sort)):
            if (sort[i] in lili) or (sort[i][1] < 0.01):
                continue
            else:
                lili.append(sort[i])
        return lili

    def rmmult_noarr(sort):
        lili = []
        for i in range(len(sort)):
            if (sort[i][0] in lili) or (sort[i][1] == 0):
                continue
            else:
                lili.append(sort[i][0])
        return lili

    tfidf_res_li = []
    num = 0

    while num < len(v_li):
        oldcorpus = v_li[num]

        stopwords = ['가운데', '대체로', '간의', '크게', '같', '아니', '있', '되', '두', '받', '많', '크', '좋', '따르', '만들', '시키', '그러',
                     '하나',
                     '모르', '데', '자신', '어떤', '명', '앞', '번', '보이', '나', '어떻', '월', '들', '이렇', '점', '싶', '좀', '잘', '통하',
                     '놓',
                     '란', '이나', '것', '이', '그', '수', '등', '의', '때', '경우', '및', '를', '대한', '사용', '위', '때문', '약', '제', '후',
                     '다른', '중', '이상', '일', '은', '위해', '가지', '시', '로', '세', '과', '다음', '두', '더', '또한', '함', '하나', '우리',
                     '개',
                     '관', '속', '가장', '모두', '점', '곳', '전', '또', '통해', '여러', '대해', '안', '즉', '모든', '내', '뒤', '보고', '이후',
                     '데',
                     '로서', '다시', '관련', '알', '비', '그것', '나', '일반', '각', '뜻', '정도', '사이', '사실', '도', '따라서', '고', '못', '예',
                     '간', '동안', '매우', '정', '지금', '증', '임', '공', '해', '음', '앞', '번', '볼', '여기', '거나', '자', '부', '인', '날',
                     '바로', '대부분', '기', '바', '주로', '뿐', '직접', '부터', '계속', '일부', '주', '재', '아래', '더욱', '초', '각각', '동시',
                     '권',
                     '년', '곧', '온', '거의', '먼저', '비롯', '역시', '몇', '반드시', '거', '보', '단', '듯', '가장', '서로', '모든', '에 대한 ',
                     ' 수 ', '모두', '의', '에', '대해', '그런데', '으로', '이것', '대로', '것', '그대로', '그리고', '것도', '수가', '이 ', '해야',
                     '지금',
                     '할 수가', '할 수', '각종', '요게', '여기', '혹시', '우리', '한번', '이번', '당신', '이렇게', '차고', '어떻게', '뭐', '깜짝', '거지',
                     '싶어서', '그래서', '정말', '이런', '도저히', '거죠', '하면은', '이제', '그렇게', '그럼', '많이', '이거', '그거', '저거', '누가',
                     '그래',
                     '그냥', '바로', '누가', '다시', '그래도', '간단히', '거야', '이따', '00', '근데', '결국', '이때', '누가', '그런', '딱', '일단',
                     '보면',
                     '하나', '어디', '부터', '원', '위하', '나오', '중', '못하', '그렇', '오', '대하', '한', '지', '하']
        corpus = []
        ########키워드 추출#########
        vectorizer = TfidfVectorizer()
        sp_matrix = vectorizer.fit_transform(oldcorpus)

        word2id = defaultdict(lambda: 0)
        for idx, feature in enumerate(vectorizer.get_feature_names()):
            word2id[feature] = idx

        li = []

        for i, sentence in enumerate(oldcorpus):
            morph = twitter.nouns(sentence)  # 명사만 사용한다.

            a = [(token, sp_matrix[i, word2id[token]]) for token in morph]  # 명사에서만 추출 #morph
            sortt = sorted(a, key=lambda k: k[1])
            sortt = rmmult(sortt)

            li += [(token, sp_matrix[i, word2id[token]]) for token in morph]  # 전체 단어 li에 저장

            sort_li = sorted(li, key=lambda k: k[1])

        a = sort_li[::-1]  # 문서 전체 tfidf 순위, # a = sort_li[-50:][::-1] --> 50개

        lili = rmmult_noarr(a)
        good_lili = []

        for j in lili:
            if j not in stopwords:
                good_lili.append(j)
            else:
                continue

        tfidf_res_li.append(good_lili[:5])  # 상위 10개
        num += 1

    ##################문장 요약 추출#################

    summary_li = []

    num = 0
    while num < len(v_li):
        str = ''.join(v_li[num])
        # print(str)

        lexrank = LexRank()
        lexrank.summarize(str)

        summaries = lexrank.probe(2)  # 1줄 요약
        summary_li.append(summaries)
        num += 1

    f = open(text_path, 'r', encoding='utf-8')
    dataset = f.read()

    sen_tag = []  # tagging job
    verb_1 = ['다', '요', '죠', '까', '니']

    dataset = f.read()
    morph = twitter.pos(dataset)
    sen_tag.append(morph)
    verb = []
    for s1 in sen_tag:
        for word, tag in s1:
            if tag == 'Verb':
                if word[-1] in verb_1:
                    verb.append(word)
    verb = list(set(verb))
    for i in verb:
        dataset = dataset.replace(i, i + '.\n')

    # 형태소 분석기로 부족한 동사들 dotword로 충전
    dotword = ['립니다', '까요', '니다', '었다', '해요', '혔다', '였다', '든요', '거든요', '아요', '습니다', '합니다', '입니다', '랍니다', '씁시다', '합시다',
               '하죠', '보죠', '이죠', '렇죠', '시죠', '었고요', '였고요', '니고요', '렇고요', '니까요', '는데>요', '니까', '이에요', '잖아요', "고요", "구요",
               "게요"]
    for i in dotword:
        dataset = dataset.replace(i + ' ', i + '.\n')

    dataset2 = dataset.split("\n")

    keywords = []
    for x in tfidf_res_li:
        for j in x:
            keywords.append(j)
    #print(keywords)

    # lexrank부분(test)
    sens = []
    for x in summary_li:
        for j in x:
            sens.append(j)
    #print(sens)
    f.close()

    document = Document()
    document.add_heading('원본.txt', 0)
    # styles 변경
    style = document.styles['Normal']
    style.font.size = Pt(11)
    for j in dataset2:
        document.add_paragraph(j)

    path = "../../../stt/origin_" + filename + ".docx"
    document.save(path)
    document = Document(path)

    # 하이라이트
    for paragraph in document.paragraphs:
        for j in keywords:
            start = paragraph.text.find(j)
            if start > -1:
                pre = paragraph.text[:start]
                # post = paragraph.text[start + len(j)]
                last = paragraph.text[start + len(j):]
                paragraph.text = pre
                paragraph.add_run(j)
                paragraph.runs[1].font.highlight_color = WD_COLOR_INDEX.YELLOW
                paragraph.add_run(last)

    for paragraph in document.paragraphs:
        for sen in sens:
            if sen in paragraph.text:
                print(sen)
                start2 = paragraph.text.find(sen)
                pre2 = paragraph.text[:start2]
                post2 = paragraph.text[len(sen):]
                paragraph.text = pre2
                paragraph.add_run(sen)
                paragraph.runs[1].font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    document.save(path)
    print("원본 텍스트:", path)

    dictkey = list(dict.keys())
    for i in range(len(dictkey)):
        if dictkey[i] == 0:
            dictkey[i] = '기술공학'
        elif dictkey[i] == 1:
            dictkey[i] = '문화예술'
        elif dictkey[i] == 2:
            dictkey[i] = '인문과학'
        elif dictkey[i] == 3:
            dictkey[i] = '역사문화'
        elif dictkey[i] == 4:
            dictkey[i] = '사회과학'
        elif dictkey[i] == 5:
            dictkey[i] = '생활취미스포츠'
        else:
            dictkey[i] = '자연과학'

    document2 = Document()

    import docx

    document2.add_heading('요약본', 0)

    for i in range(len(tfidf_res_li)):
        document2.add_heading('{}'.format(i + 1), 1)
        # for i in range(len(dictkey)):
        document2.add_heading('키워드', 3)
        kk = ', '.join(tfidf_res_li[i])
        print(kk)
        document2.add_paragraph(kk)
        document2.add_heading('내용요약', 3)
        ll = ' '.join(summary_li[i])
        document2.add_paragraph(ll)

    path = "../../../summary/smry_" + filename + ".docx"
    document2.save(path)

    print("요약 텍스트:", path)


if __name__ == '__main__':
    text_path = sys.argv[1]
    filename = text_path.split('/')[-1].split('.')[0]

